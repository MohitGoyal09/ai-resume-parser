import pypdfium2 as pdfium
import docx 
import io
import logging 

logger = logging.getLogger(__name__)

# Extract text from a PDF file
def extract_text_from_pdf(file_content: bytes) -> str:
    text_content = ""
    try:
        pdf = pdfium.PdfDocument(file_content)
        for i in range(len(pdf)):
            page = pdf[i]
            textpage = page.get_textpage()
            text_content += textpage.get_text_range() + "\n\n"
            textpage.close()
            page.close()
        pdf.close()
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}", exc_info=True)
        return ""
    return text_content.strip()

# Extract text from a DOCX file
def extract_text_from_docx(file_content: bytes) -> str:
    text_content = ""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        for para in doc.paragraphs:
            text_content += para.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}", exc_info=True)
        return ""
    return text_content.strip()

# Extract text from a resume file   
def extract_text_from_resume(filename: str, file_content: bytes) -> str:
    if not filename:
        raise ValueError("Filename cannot be empty.")
    
    file_ext = filename.split('.')[-1].lower()

    if file_ext == "pdf":
        return extract_text_from_pdf(file_content)
    elif file_ext == "docx":
        return extract_text_from_docx(file_content)
    elif file_ext == "txt":
         try:
            return file_content.decode('utf-8', errors='strict').strip()
         except UnicodeDecodeError:
            logger.warning(f"Could not decode {filename} as UTF-8, trying common encodings.")
            # Trying other common encodings if UTF-8 fails
            for encoding in ['latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding, errors='strict').strip()
                except UnicodeDecodeError:
                    continue
            logger.error(f"Failed to decode {filename} with common encodings.")
            raise ValueError(f"Unsupported text encoding for .txt file: {filename}")
    else:
        raise ValueError(f"Unsupported file type: .{file_ext}. Please upload PDF, DOCX, or TXT.")